"""Stage 1: turn source files into normalized Documents.

Job postings come from a committed snapshot, which is the demo path:
instant, reproducible, and offline-safe. Syllabi are parsed from PDF or
DOCX at upload time.

Nothing in this module raises on bad input. A document that cannot be
parsed is logged and skipped, because one malformed PDF must not take down
a batch ingest.
"""

import logging
import uuid
from pathlib import Path

from app.contracts import Document, DocumentKind

logger = logging.getLogger(__name__)

SNAPSHOT_PATH = Path("/app/data/jobs_snapshot.json")

# Below this, a "parsed" document is almost certainly a failed text-layer
# extraction rather than a real syllabus.
MIN_USABLE_CHARS = 200


def load_job_snapshot(path: Path | None = None) -> list[Document]:
    """Load the committed job posting corpus.

    Extra fields in the snapshot (company, source_url, seniority) are kept
    out of Document deliberately: they belong to the JobPosting node in the
    graph, not to the pipeline contract.
    """
    import json

    snapshot = path or SNAPSHOT_PATH
    if not snapshot.exists():
        logger.warning("Job snapshot not found at %s", snapshot)
        return []

    raw = json.loads(snapshot.read_text(encoding="utf-8"))
    documents = []
    for item in raw:
        try:
            documents.append(
                Document(
                    doc_id=item["doc_id"],
                    kind=DocumentKind.JOB_POSTING,
                    title=item["title"],
                    raw_text=item["raw_text"],
                    source=item.get("source", "unknown"),
                    posted_date=item.get("posted_date"),
                )
            )
        except Exception as exc:
            logger.warning("Skipping malformed posting %s: %s", item.get("doc_id"), exc)
    return documents


def load_job_metadata(path: Path | None = None) -> dict[str, dict]:
    """Snapshot fields that belong on the graph node rather than the Document."""
    import json

    snapshot = path or SNAPSHOT_PATH
    if not snapshot.exists():
        return {}

    raw = json.loads(snapshot.read_text(encoding="utf-8"))
    return {
        item["doc_id"]: {
            "company": item.get("company", ""),
            "source_url": item.get("source_url", ""),
            "seniority": item.get("seniority", "mid"),
        }
        for item in raw
    }


def parse_syllabus(path: Path) -> Document | None:
    """Parse one syllabus file. Returns None rather than raising.

    PDF text extraction is the single most failure-prone step in the whole
    pipeline, so every failure mode here degrades to a skip.
    """
    try:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = _extract_pdf(path)
        elif suffix in {".docx", ".doc"}:
            text = _extract_docx(path)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.warning("Failed to parse %s: %s", path.name, exc)
        return None

    if len(text.strip()) < MIN_USABLE_CHARS:
        logger.warning(
            "Text layer too thin in %s (%d chars); flagged for skip",
            path.name,
            len(text.strip()),
        )
        return None

    return Document(
        doc_id=f"syl-{uuid.uuid4().hex[:8]}",
        kind=DocumentKind.SYLLABUS,
        title=path.stem.replace("_", " ").replace("-", " ").strip(),
        raw_text=text,
        source=path.name,
    )


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    return "\n".join(p.text for p in DocxDocument(str(path)).paragraphs)


def load_syllabi(directory: Path) -> list[Document]:
    """Parse every file in a directory, skipping whatever fails."""
    if not directory.exists():
        logger.warning("Syllabus directory not found: %s", directory)
        return []

    documents = []
    for path in sorted(directory.iterdir()):
        if path.is_file() and not path.name.startswith("."):
            parsed = parse_syllabus(path)
            if parsed is not None:
                documents.append(parsed)
    return documents
