"""Fill the TetraTHON 2026 Statement of Work.

Reproduces the section order, headings and field layout of the supplied PDF
exactly, so a judge comparing the two sees the same document with the blanks
completed. Figures come from the running system rather than being written
from memory.
"""

import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = pathlib.Path("docs/TETRA030_Statement_of_Work.docx")

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTE = RGBColor(0x66, 0x66, 0x66)
RULE = "999999"


def shade(cell, hex_fill):
    element = OxmlElement("w:shd")
    element.set(qn("w:val"), "clear")
    element.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(element)


def borders(table):
    tbl_pr = table._tbl.tblPr
    marks = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "4")
        line.set(qn("w:color"), RULE)
        marks.append(line)
    tbl_pr.append(marks)


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    return p


def hint(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTE
    return p


def body(doc, text, space_after=6, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK
    return p


def fields(doc, rows):
    """The label/value pairs the PDF renders as a two-column form."""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders(table)
    for label, value in rows:
        cells = table.add_row().cells
        shade(cells[0], "F2F1EA")

        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = INK

        run = cells[1].paragraphs[0].add_run(value)
        run.font.size = Pt(9)
        run.font.color.rgb = INK

        cells[0].width = Pt(150)
        cells[1].width = Pt(310)
    return table


def main() -> int:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)

    for section in doc.sections:
        section.top_margin = Pt(46)
        section.bottom_margin = Pt(46)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ---- masthead ----------------------------------------------------
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("TetraTHON 2026")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = INK

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    run = sub.add_run(
        "An Indo-French AI Innovation Sprint · "
        "Navrachana Innovation Foundation (NIF)"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = MUTE

    banner = doc.add_paragraph()
    banner.paragraph_format.space_after = Pt(4)
    run = banner.add_run("PROJECT STATEMENT OF WORK")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = INK

    # ---- team and submission -----------------------------------------
    heading(doc, "TEAM & SUBMISSION DETAILS")
    fields(
        doc,
        [
            ("Team Name", "Debugging Demons (Team ID: TETRA030)"),
            ("Team Lead Name", "Harshul Rathod"),
            ("Team Lead Email", "harshulrathod1640@gmail.com"),
            ("Team Lead Phone", ""),
            ("Date Submitted", "2 August 2026"),
        ],
    )

    # ---- sector track -------------------------------------------------
    heading(doc, "SECTOR TRACK")
    track = doc.add_paragraph()
    track.paragraph_format.space_after = Pt(4)
    for label, chosen in (
        ("HealthTech", False),
        ("FinTech", False),
        ("AgriTech", False),
        ("EdTech", True),
    ):
        run = track.add_run(f"{'☒' if chosen else '☐'} {label}    ")
        run.font.size = Pt(10)
        run.bold = chosen
        run.font.color.rgb = INK if chosen else MUTE

    # ---- project information -----------------------------------------
    heading(doc, "PROJECT INFORMATION")
    fields(
        doc,
        [
            ("Project / Solution Name", "Vedha"),
            (
                "Problem Statement No.",
                "Track D, Problem Statement 1 — Dynamic Syllabus and "
                "Industry Skill-Gap Synchronizer",
            ),
            ("Industry Partner (if applicable)", "None"),
        ],
    )

    # ---- assumptions ---------------------------------------------------
    heading(doc, "ASSUMPTIONS")
    hint(
        doc,
        "List any assumptions made about data availability, tools, APIs, or "
        "constraints.",
    )
    for line in [
        "Job market data comes from free public APIs (Arbeitnow and Remotive). "
        "These serve current listings only, so the corpus spans roughly one "
        "month. Month-granularity trend analysis is therefore supportable; a "
        "multi-year forecast is not.",
        "That corpus is predominantly software roles. Where a discipline is "
        "under-represented, the system reports insufficient evidence rather "
        "than issuing a verdict, so this is handled as a stated limitation "
        "rather than an unstated one.",
        "Syllabi are published university PDFs with an extractable text "
        "layer. Scanned documents without one are rejected rather than passed "
        "through OCR, since silently misread text would corrupt the audit.",
        "Skill extraction is constrained to a curated taxonomy. A language "
        "model naming skills freely produces plausible entries that no job "
        "posting uses, so the vocabulary is fixed and the model may only "
        "select within it.",
        "Embeddings run locally (sentence-transformers, all-MiniLM-L6-v2), so "
        "no per-request API cost is incurred for matching.",
        "Judging is on a laptop with Docker available. The system runs "
        "entirely on the local machine with no external services beyond the "
        "two job APIs and one language model endpoint.",
    ]:
        body(doc, line, bullet=True)

    doc.add_page_break()

    # ---- scope ---------------------------------------------------------
    heading(doc, "SCOPE OF WORK")
    hint(
        doc,
        "Describe what the team will build during the 36-hour sprint, "
        "including key features and boundaries.",
    )
    body(
        doc,
        "Vedha audits a university syllabus against live job market data and "
        "reports not only which skills are missing but what it would cost to "
        "teach them. A six-stage pipeline (Ingest, Extract, Embed, Match, "
        "Score, Augment) parses syllabus documents, resolves learning "
        "outcomes to a curated skill taxonomy, compares them against skills "
        "demanded by real job postings, and returns an evidence-backed report.",
    )
    body(doc, "Delivered within the sprint:")
    for line in [
        "Syllabus ingestion from PDF, Word, Markdown and plain text, "
        "including automatic splitting of multi-course curriculum documents.",
        "A skill ontology of 487 skills, 1,646 aliases and 626 prerequisite "
        "relationships held in Neo4j, spanning computing, civil, mechanical, "
        "electrical and process engineering.",
        "Gap analysis scoring each course against demand in its own subject "
        "area rather than the whole market, with every finding citing the "
        "job postings behind it.",
        "Prerequisite-aware cost estimation: the graph answers how many "
        "teaching steps separate a course from a missing skill, which turns "
        "“you lack X” into “X costs two prerequisites”.",
        "A teaching sequence produced by topological layering of the gap "
        "subgraph, so nothing is scheduled before its groundwork.",
        "Proposed syllabus modifications that extend existing units rather "
        "than requiring a redesign.",
        "An NBA-style accreditation view mapping course outcomes to the "
        "twelve Programme Outcomes.",
        "A web interface with an interactive ontology graph and an upload "
        "path so any syllabus can be analysed live.",
    ]:
        body(doc, line, bullet=True)

    body(doc, "Explicitly out of scope for the sprint:")
    for line in [
        "OCR for scanned documents without a text layer.",
        "Multi-year demand forecasting, which the available data cannot "
        "support.",
        "Authentication, multi-tenancy and institutional data governance.",
        "Automatic writing of approved changes back into university systems; "
        "the tool produces proposals for a curriculum committee to adopt.",
    ]:
        body(doc, line, bullet=True)

    # ---- goals ----------------------------------------------------------
    heading(doc, "PROJECT GOALS")
    for line in [
        "Replace a manual curriculum review that takes months with an audit "
        "that runs in minutes and cites its evidence.",
        "Move beyond keyword comparison. Recording which skills are "
        "prerequisites for which allows the system to answer what closing a "
        "gap would cost, which is the question a curriculum committee "
        "actually faces.",
        "Work across engineering disciplines rather than computing alone, so "
        "the tool serves a whole institution.",
        "Report honestly. Where the evidence is too thin to judge a subject "
        "area, say so and withhold both the findings and the score rather "
        "than presenting a confident number built on nothing.",
        "Keep every claim checkable: each figure traces back to named job "
        "postings a reader can open.",
    ]:
        body(doc, line, bullet=True)

    # ---- impact ---------------------------------------------------------
    heading(doc, "IMPACT & SCALABILITY")
    hint(
        doc,
        "How does this solution support inclusive growth and scale beyond the "
        "sprint? (theme: AI for Inclusive Growth and Scalable Impact)",
    )
    for line in [
        "Curriculum lag hits hardest at institutions without industry "
        "advisory boards or dedicated curriculum staff. Well-resourced "
        "universities already run employer consultations; smaller colleges "
        "cannot. Automating the analysis puts the same evidence in the hands "
        "of institutions that could not otherwise afford it, which is where "
        "the inclusive-growth argument sits.",
        "Students bear the cost of the lag through weaker employability. "
        "Shortening the revision cycle from years to a term reaches every "
        "student in a programme, not only those who can supplement their "
        "degree independently.",
        "Coverage across civil, mechanical, electrical and process "
        "engineering matters for the same reason: non-computing disciplines "
        "are usually last to receive tooling of this kind.",
        "The architecture scales by substitution rather than rewrite. Each "
        "pipeline stage has one responsibility behind a stable interface, so "
        "a richer job source, a different embedding model or a larger "
        "taxonomy can be swapped in without touching its neighbours.",
        "Cost scales sub-linearly. Embeddings run locally at no per-request "
        "cost, and language model responses are cached by content hash, so "
        "re-analysing an unchanged syllabus is free.",
        "The taxonomy is data, not code. Extending the system to a new "
        "discipline means adding entries to a JSON file that a subject "
        "expert can review, not modifying the pipeline.",
    ]:
        body(doc, line, bullet=True)

    # ---- deliverables ----------------------------------------------------
    heading(doc, "PROJECT DELIVERABLES")
    hint(
        doc,
        "e.g. working prototype, demo video, pitch deck, source code "
        "repository",
    )
    for line in [
        "Working prototype: a running full-stack application (FastAPI "
        "backend, Next.js frontend, Neo4j and ChromaDB), reproducible with a "
        "single Docker Compose command.",
        "Source code repository: github.com/Harshul1484/TETRA030, developed "
        "throughout on feature branches and pull requests.",
        "Seeded dataset: published syllabi from The Maharaja Sayajirao "
        "University of Baroda and NIT Tiruchirappalli, audited against a live "
        "corpus of real job postings.",
        "Automated test suite of 164 tests covering the pipeline, scoring and "
        "API contract.",
        "Pitch deck for the Round B presentation.",
        "Documentation: README with architecture diagrams, plus this "
        "Statement of Work.",
    ]:
        body(doc, line, bullet=True)

    doc.add_page_break()

    # ---- team ------------------------------------------------------------
    heading(doc, "TEAM STRUCTURE")
    table = doc.add_table(rows=1, cols=4)
    borders(table)
    for index, label in enumerate(
        ["Name", "Role on Project", "Institution", "Email"]
    ):
        cell = table.rows[0].cells[index]
        shade(cell, "F2F1EA")
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)

    for name, role, institution, email in [
        (
            "Harshul Rathod",
            "Team Lead — architecture, backend, frontend, ML pipeline",
            "The Maharaja Sayajirao University of Baroda",
            "harshulrathod1640@gmail.com",
        ),
        (
            "Manav Vadgama",
            "Research and documentation",
            "The Maharaja Sayajirao University of Baroda",
            "manavvadgama28@gmail.com",
        ),
        (
            "Huzefa Vohra",
            "Testing and presentation",
            "The Maharaja Sayajirao University of Baroda",
            "huzialtair@gmail.com",
        ),
    ]:
        cells = table.add_row().cells
        for index, value in enumerate([name, role, institution, email]):
            run = cells[index].paragraphs[0].add_run(value)
            run.font.size = Pt(8.5)

    # ---- stack -----------------------------------------------------------
    heading(doc, "PROPOSED TECH STACK")
    fields(
        doc,
        [
            ("Backend", "Python 3.11, FastAPI, Pydantic"),
            ("Frontend", "Next.js 15, TypeScript, React, Tailwind CSS"),
            (
                "Graph database",
                "Neo4j 5.26 — holds the skill ontology; prerequisite "
                "depth is a variable-length path search, which a relational "
                "schema answers only with a recursive query that degrades "
                "with every hop",
            ),
            (
                "Vector store",
                "ChromaDB (embedded) with sentence-transformers "
                "all-MiniLM-L6-v2, run locally so matching costs nothing per "
                "request",
            ),
            (
                "Language model",
                "Claude API for outcome extraction and modification "
                "proposals, constrained to the curated taxonomy, with a "
                "content-hash disk cache and a three-level fallback chain",
            ),
            ("Document parsing", "pdfplumber, python-docx"),
            ("Job market data", "Arbeitnow and Remotive public APIs"),
            ("Infrastructure", "Docker Compose"),
            (
                "Quality",
                "pytest (164 tests), OpenAPI-generated TypeScript types to "
                "prevent contract drift between backend and frontend",
            ),
        ],
    )

    # ---- cost ------------------------------------------------------------
    heading(doc, "FUTURE COST TO SCALE")
    hint(
        doc,
        "If this project were taken beyond the hackathon into a prototype or "
        "startup, rough order-of-magnitude cost to build and launch (not "
        "required to be precise).",
    )
    body(
        doc,
        "Order-of-magnitude estimate for a twelve-month path from prototype "
        "to a product serving a first cohort of institutions.",
    )

    table = doc.add_table(rows=1, cols=3)
    borders(table)
    for index, label in enumerate(["Item", "Annual estimate (INR)", "Note"]):
        cell = table.rows[0].cells[index]
        shade(cell, "F2F1EA")
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)

    for item, amount, note in [
        (
            "Engineering (2 developers)",
            "18,00,000 – 24,00,000",
            "The dominant cost. Covers hardening, institutional integrations "
            "and support.",
        ),
        (
            "Licensed job market data",
            "3,00,000 – 6,00,000",
            "The single largest technical constraint today. Free APIs are "
            "software-skewed and hold one month of history; a licensed feed "
            "removes the limitation this document records under Assumptions.",
        ),
        (
            "Cloud infrastructure",
            "1,20,000 – 2,40,000",
            "Managed Neo4j, application hosting, object storage. Scales with "
            "institutions, not with students.",
        ),
        (
            "Language model usage",
            "60,000 – 1,50,000",
            "Held down by caching and by running embeddings locally. Cost "
            "recurs only when a syllabus actually changes.",
        ),
        (
            "Curriculum and subject expertise",
            "2,00,000 – 4,00,000",
            "Part-time review of the taxonomy as it extends into further "
            "disciplines.",
        ),
        (
            "Compliance and accreditation alignment",
            "1,00,000 – 2,00,000",
            "NBA and NAAC reporting formats, data handling agreements.",
        ),
        (
            "Total",
            "25,80,000 – 39,90,000",
            "Roughly USD 31,000 – 48,000 for year one.",
        ),
    ]:
        cells = table.add_row().cells
        for index, value in enumerate([item, amount, note]):
            run = cells[index].paragraphs[0].add_run(value)
            run.font.size = Pt(8.5)
            if item == "Total":
                run.bold = True

    body(
        doc,
        "Marginal cost per additional institution is low: the taxonomy, "
        "ontology and job corpus are shared, so onboarding adds syllabus "
        "storage and support rather than new infrastructure.",
        space_after=4,
    )

    # ---- acknowledgement --------------------------------------------------
    heading(doc, "ACKNOWLEDGEMENT")
    body(
        doc,
        "By submitting this Statement of Work, the team confirms the "
        "information above is accurate and agrees to the TetraTHON 2026 "
        "judging criteria and event guidelines.",
    )

    fields(
        doc,
        [
            ("Team Lead Signature", ""),
            ("Date", "2 August 2026"),
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(18)
    run = footer.add_run(
        "TetraTHON 2026 · 31 July – 2 August 2026 · "
        "Navrachana University, Vadodara · nif@navrachana.edu.in"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = MUTE

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
